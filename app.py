# Disable TensorFlow oneDNN and deprecation logs
import os
os.environ["TF_CPP_MIN_LOG_LEVEL"] = "3"  # Suppress TF logs
os.environ["TF_ENABLE_ONEDNN_OPTS"] = "0"  # Disable oneDNN logs
from flask import Flask, render_template, request, redirect, flash, url_for
import sqlite3
import re

from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.serving import WSGIRequestHandler
WSGIRequestHandler.protocol_version = "HTTP/1.1"  # Keeps connections alive

# LANGUAGE TRANSLATOR LIBRARIES
from googletrans import Translator, LANGUAGES
import fitz  # PyMuPDF for PDFs
from docx import Document  # We use this now for .docx generation
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image
from flask import Flask, request, jsonify, send_file, flash
from gtts import gTTS
import os
import uuid

translator = Translator()

# FAKE NEWS DETECTION LIBRARIES
import pickle
import re
import string

# CONTENT PLAGIARISM DETECTION LIBRARIES
import requests
import bs4
import spacy
import PyPDF2
from flask import Flask, render_template, request, jsonify
from googlesearch import search
from summa import summarizer
from sentence_transformers import SentenceTransformer, util
from werkzeug.utils import secure_filename


# TEXT SUMMARIZATION LIBRARIES
import torch
import nltk
import logging
from transformers import PegasusTokenizer, PegasusForConditionalGeneration
from werkzeug.utils import secure_filename

# SENTIMENT ANALYSIS MODEL
import pandas as pd
import torch
from transformers import AutoTokenizer, AutoModelForSequenceClassification
import io
from fpdf import FPDF
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import seaborn as sns
from wordcloud import WordCloud
import tempfile
import numpy as np



# Append the path to the nltk_data folder in your Flask app
nltk.data.path.append("./nltk_data")
# Now load 'punkt' without downloading
nltk.data.load('tokenizers/punkt/english.pickle')
print("NLTK 'punkt' loaded successfully!")


# Configure Logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

# SENTIMENT ANALYSIS LIBRARIES

def init_db():
    conn = sqlite3.connect('feedbacks.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS feedbacks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT NOT NULL,
            phone TEXT NOT NULL,
            feedback TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

app = Flask(__name__)
init_db()

app.secret_key = 'analytics_hub_secret_key_2025'


app.config['TIMEOUT'] = 60  # Increase request timeout to 60 seconds
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # Disables caching

nlp = spacy.load("en_core_web_sm")
local_model_path = r"C:\Users\nihar\Desktop\8th Sem\AI Tools\all-MiniLM-L6-v2"
model = SentenceTransformer(local_model_path)

# Set plagiarism_upload folder
UPLOAD_FOLDER = "plagiarism_uploads"
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
ALLOWED_EXTENSIONS = {"pdf"}

UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

AUDIO_FOLDER = "static/audio"
os.makedirs(AUDIO_FOLDER, exist_ok=True)  # Ensure the folder exists

@app.route('/')
def home():
    return render_template('index2.html')

@app.route('/about')
def about():
    return render_template('About.html')  # Example, replace with about.html if needed

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    if request.method == 'POST':
        email = request.form.get('email')
        phone = request.form.get('phone')
        feedback = request.form.get('feedback')

        # Server-side Validations
        if not re.match(r'^\S+@\S+\.\S+$', email):
            flash('Invalid email address!', 'danger')
            return redirect(url_for('contact'))
        
        if not re.match(r'^[0-9]{10}$', phone):
            flash('Phone number must be exactly 10 digits!', 'danger')
            return redirect(url_for('contact'))

        # Save into Database
        conn = sqlite3.connect('feedbacks.db')
        c = conn.cursor()
        c.execute('INSERT INTO feedbacks (email, phone, feedback) VALUES (?, ?, ?)', (email, phone, feedback))
        conn.commit()
        conn.close()

        flash('Thank you for your feedback!', 'success')
        return redirect(url_for('contact'))

    return render_template('contact.html')

@app.route('/view-feedbacks')
def view_feedbacks():
    conn = sqlite3.connect('feedbacks.db')
    c = conn.cursor()
    c.execute('SELECT * FROM feedbacks')
    rows = c.fetchall()
    conn.close()
    return render_template('view_feedbacks.html', feedbacks=rows)


@app.route('/privacy-policy')
def privacy_policy():
    return render_template('Privacy_policy.html')

# LANGUAGE TRANSLATOR
# Function to extract text from uploaded files
def doc2text(file_path):
    if file_path.endswith('.pdf'):
        with fitz.open(file_path) as doc:
            text = ''.join(page.get_text() for page in doc)
        return text
    elif file_path.endswith(('.docx', '.DOCX', '.Docx')):
        document = Document(file_path)
        return '\n'.join(paragraph.text for paragraph in document.paragraphs)
    elif file_path.endswith(('.html', '.HTML')):
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'lxml')
        return soup.get_text(separator='\n', strip=True)
    elif file_path.endswith(('.jpg', '.jpeg', '.png')):
        return pytesseract.image_to_string(Image.open(file_path))
    return None

# Route for Translator Page
@app.route('/translator')
def translator_page():
    return render_template('Translator.html', languages=LANGUAGES)

# API Route for Text Translation
@app.route('/translate', methods=['POST'])
def translator_2():
    languages = LANGUAGES
    text = request.form.get('text')
    language = request.form.get('language', 'en')

    if not text:
        return jsonify({'error': 'Please enter text to translate'}), 400

    if len(text) > 15000:
        return jsonify({'error': 'Character limit exceeded (Max: 15,000)'}), 400
    print(text)
    print(language)
    # print(type(translator))
    print(len(text))
    try:
        translated_text = translator.translate(text, dest=language).text

        # Generate unique filename
        audio_filename = f"{uuid.uuid4()}.mp3"
        audio_path = os.path.join(AUDIO_FOLDER, audio_filename)

        # Convert text to speech
        tts = gTTS(translated_text, lang=language)
        tts.save(audio_path)

        return jsonify({
            'translated_text': translated_text,
            'audio_url': f"/static/audio/{audio_filename}"
        })
    except Exception as e:
        return jsonify({'error': f'Translation failed: {str(e)}'}), 500
    
    

# API Route for File Upload Translation


from docx import Document

def generate_docx(translated_text, docx_path):
    doc = Document()
    doc.add_heading('Translated Document', level=1)

    # Split into paragraphs if needed
    for paragraph in translated_text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    doc.save(docx_path)


@app.route('/translate_file', methods=['POST'])
def translate_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    language = request.form.get('language')

    if not file or file.filename == '':
        return jsonify({'error': 'Please select a file to translate'}), 400

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)

    extracted_text = doc2text(file_path)  # Extract text from file

    if not extracted_text:
        return jsonify({'error': 'Unsupported file format'}), 400

    if len(extracted_text) > 15000:
        return jsonify({'error': 'Character limit exceeded (Max: 15,000)'}), 400

    translated_text = translator.translate(extracted_text, dest=language).text
    print(translated_text)

    try:
        # Define DOCX file path
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], 'translated.docx')

        # Generate the DOCX instead of PDF
        generate_docx(translated_text, docx_path)

        return jsonify({'download_url': '/download_translated_docx'})

    except Exception as e:
        return jsonify({'error': f"Failed to generate DOCX: {str(e)}"}), 500

@app.route('/download_translated_docx')
def download_translated_docx():
    try:
        return send_file(
            os.path.join(app.config['UPLOAD_FOLDER'], 'translated.docx'), 
            as_attachment=True,
            download_name='translated.docx'  # Set download filename nicely
        )
    except Exception as e:
        return jsonify({'error': f"Failed to download DOCX: {str(e)}"}), 500




# FAKE NEWS DETECTION
# Load the trained model and vectorizer
model_path = 'static/FakeNewsmodel.pkl'
vectorizer_path = 'static/FakeNewsvectorizer.pkl'

with open(model_path, 'rb') as model_file:
    classifier = pickle.load(model_file)

with open(vectorizer_path, 'rb') as vectorizer_file:
    vectorizer = pickle.load(vectorizer_file)

# Text preprocessing function
def preprocess(text):  
    text = text.lower()
    text = re.sub(r'\[.*?\]', '', text)  
    text = re.sub(r'https?://\S+|www\.\S+', '', text)  
    text = re.sub(r'<.*?>+', '', text)  
    text = re.sub(r'[%s]' % re.escape(string.punctuation), '', text)  
    text = re.sub(r'\n', ' ', text)  
    text = re.sub(r'\w*\d\w*', '', text)  
    text = re.sub(r'\W', ' ', text)  
    return text.strip()

# Route for Fake News Detection
@app.route("/fake-news", methods=['GET', 'POST'])
def fake_news():
    if request.method == "POST":
        user_input = request.form.get("news_text", "").strip()
        if not user_input:
            return jsonify({"error": "Please enter news text to check."})
        
        processed_text = preprocess(user_input)
        user_input_vectorized = vectorizer.transform([processed_text])
        prediction = classifier.predict(user_input_vectorized)
        result = "Fake News" if prediction[0] == 1 else "Real News"
        return jsonify({"result": result})

    return render_template("FakeNews.html")


# CONTENT PLAGIARISM DETECTOR

# def allowed_file(filename):
#     return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# def extract_text_from_pdf(pdf_path):
#     """Extracts text from a PDF file."""
#     try:
#         with open(pdf_path, "rb") as file:
#             reader = PyPDF2.PdfReader(file)
#             text = ""
#             for page in reader.pages:
#                 text += page.extract_text() + "\n"
#         return text.strip()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         return ""

def preprocess_text(text):
    """Cleans text by removing special characters and extra spaces."""
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s.!?]', '', text)
    return text

def extract_key_sentences_textrank(text, num_sentences=3):
    """Extracts important sentences using TextRank."""
    summary = summarizer.summarize(text, ratio=0.3)
    if not summary.strip():
        return text
    doc = nlp(summary)
    sentences = [sent.text.strip() for sent in doc.sents]
    return " ".join(sentences[:num_sentences])

def extract_key_sentences_embeddings(text, num_sentences=3):
    """Finds unique & important sentences using embeddings."""
    doc = nlp(text)
    sentences = [sent.text.strip() for sent in doc.sents]
    if not sentences:
        return text
    sentence_embeddings = model.encode(sentences, convert_to_tensor=True)
    sentence_scores = sentence_embeddings.mean(dim=1).tolist()
    sorted_sentences = [sent for _, sent in sorted(zip(sentence_scores, sentences), reverse=True)]
    return " ".join(sorted_sentences[:num_sentences])

def extract_key_sentences(text):
    """Choose method based on text length."""
    if len(text.split()) < 100:
        return extract_key_sentences_textrank(text, num_sentences=3)
    return extract_key_sentences_embeddings(text, num_sentences=5)

def google_search(query, num_results=10):
    """Fetches top search results from Google."""
    results = []
    for i in search(query, num=num_results):
        results.append(i)
        if len(results) >= num_results:
            break  # 🔥 force stop after num_results
    return results

def preprocess_web_text(text):
    """Cleans web-scraped text while keeping structure."""
    text = text.lower()
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'\[.*?\]', '', text)
    text = re.sub(r'\(.*?\)', '', text)
    text = re.sub(r'[“”‘’]', '', text)
    text = re.sub(r'[^a-z0-9.,!?;:\'\"()\s]', '', text)
    return text

def scrape_website(url):
    """Scrapes text from a webpage."""
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        response = requests.get(url, headers=headers, timeout=5)
        soup = bs4.BeautifulSoup(response.text, "html.parser")
        paragraphs = soup.find_all("p")
        text = " ".join([p.get_text() for p in paragraphs])
        return preprocess_web_text(text)
    except Exception:
        return ""

def calculate_similarity(input_text, website_text):
    """Computes similarity score using SBERT."""
    if not website_text:
        return 0
    input_embedding = model.encode(input_text, convert_to_tensor=True)
    website_embedding = model.encode(website_text, convert_to_tensor=True)
    return round(util.pytorch_cos_sim(input_embedding, website_embedding).item() * 100, 2)

def check_plagiarism(input_text):
    """Checks plagiarism and returns top 3 results."""
    input_text = preprocess_text(input_text)
    key_sentences = extract_key_sentences(input_text)
    urls = google_search(key_sentences,num_results=10)
    urls = urls[:10]
    for i in urls:
        print(i)
    similarities = {}
    for url in urls:
        web_text = scrape_website(url)
        similarity = calculate_similarity(input_text, web_text)
        similarities[url] = similarity
    sorted_results = sorted(similarities.items(), key=lambda x: x[1], reverse=True)[:3]
    return sorted_results

@app.route("/plagiarism", methods=["GET", "POST"])
def plagiarism():
    if request.method == "POST":
        input_text = request.form.get("text", "").strip()

        if not input_text:
            return jsonify({"error": "No text provided"}), 400

        results = check_plagiarism(input_text)
        # Convert results from list of lists to list of dictionaries
        formatted_results = [{"url": item[0], "similarity": item[1]} for item in results]
        return jsonify(formatted_results)

    return render_template("plagiarism.html")



# TEXT SUMMARIZATION
# Load Text Summarization Model
# local_dir = r"C:\Users\nihar\Desktop\8th Sem\Hugging_Face\pegasus_model"
# tokenizer = PegasusTokenizer.from_pretrained(local_dir)
tokenizer = "c"
# summarization_model = PegasusForConditionalGeneration.from_pretrained(local_dir).to("cuda" if torch.cuda.is_available() else "cpu")
summarization_model = "xyz"

# Function to Extract Text from PDF
def extract_text_from_pdf(file):
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = "\n".join(page.get_text("text") for page in doc)
        return text.strip()
    except Exception as e:
        logging.error(f"PDF Processing Error: {e}")
        return None

# Function to Extract Text from DOCX
def extract_text_from_docx(file):
    try:
        doc = Document(file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        logging.error(f"DOCX Processing Error: {e}")
        return None

# Text Chunking to Fit Model's Token Limit
def chunk_text(text, tokenizer, max_tokens=1024):
    tokens = tokenizer.encode(text, truncation=False)
    chunks = [tokens[i:i+max_tokens] for i in range(0, len(tokens), max_tokens)]
    print(f"🔹 Total Input Tokens: {len(tokens)}")
    print(f"🔹 Number of Chunks: {len(chunks)}")
    print(f"🔹 Chunk Sizes: {[len(chunk) for chunk in chunks]}\n")
    return [tokenizer.decode(chunk, skip_special_tokens=True) for chunk in chunks], len(tokens)

# Ensure Summaries End at Proper Sentences
def ensure_complete_sentences(text):
    sentences = nltk.sent_tokenize(text)
    return " ".join(sentences[:-1]) if len(sentences) > 1 else sentences[0]

# Summarization Function
def summarize_text(text, min_ratio=0.25, max_ratio=0.30):
    try:
        device = "cuda" if torch.cuda.is_available() else "cpu"
        max_input_tokens = 1024  # Pegasus model limit
        text_chunks, total_input_tokens = chunk_text(text, tokenizer, max_tokens=max_input_tokens)

        if device == "cuda":
            torch.cuda.empty_cache()  # Free up GPU memory
            print("✅ CUDA Cache Cleared")

        summaries = []
        for idx, chunk in enumerate(text_chunks):
            input_tokens = tokenizer.encode(chunk, return_tensors="pt", truncation=True).to(device)
            input_length = input_tokens.shape[1]
            
            # Ensure the min/max summary length does not exceed safe limits
            min_summary_length = max(20, int(input_length * min_ratio))
            max_summary_length = min(512, max(min_summary_length + 30, int(input_length * max_ratio)))

            print(f"🟢 Processing Chunk {idx+1} / {len(text_chunks)}:")
            print(f"   🔸 Input Tokens: {input_length}")
            print(f"   🔸 Expected Output Tokens: {min_summary_length} - {max_summary_length}")
            
            try:
                summary_ids = summarization_model.generate(
                    input_tokens, 
                    max_length=max_summary_length, 
                    min_length=min_summary_length, 
                    length_penalty=1.5, 
                    num_return_sequences=1,
                    early_stopping=True  # Helps prevent infinite loops
                )
                output_tokens = summary_ids.shape[1]
                summary_text = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
                summary_text = ensure_complete_sentences(summary_text)
                summaries.append(summary_text)

                print(f"   🔹 Output Tokens: {output_tokens}")
                print(f"   🔹 Summary Generated ✅\n")
            
            except torch.cuda.OutOfMemoryError:
                print(f"🚨 CUDA Out of Memory on Chunk {idx+1}. Trying on CPU...")
                torch.cuda.empty_cache()
                input_tokens = input_tokens.to("cpu")
                summary_ids = summarization_model.generate(
                    input_tokens, 
                    max_length=max_summary_length, 
                    min_length=min_summary_length, 
                    length_penalty=1.5, 
                    num_return_sequences=1,
                    early_stopping=True,
                    no_repeat_ngram_size=3
                )
                output_tokens = summary_ids.shape[1]
                summary_text = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
                summary_text = ensure_complete_sentences(summary_text)
                summaries.append(summary_text)
                print(f"✅ Successfully Processed Chunk {idx+1} on CPU")
                # Ensure summaries exist before joining
        if summaries:
            final_summary = " ".join(summaries).replace("<n>", " ")
        else:
            final_summary = "No summary generated. Please check the input text."
        return final_summary

    except Exception as e:
        logging.error(f"Summarization Error: {str(e)}")  # Ensure error is logged properly
        return "An error occurred while summarizing the text."



@app.route('/summarization')
def summarization_page():
    return render_template('summarization.html')

@app.route('/summarize_text', methods=['POST'])
def summarize_text_route():
    try:
        text = request.form.get('text', '').strip()
        if not text:
            return jsonify({"error": "No text provided!"}), 400

        summary = summarize_text(text)
        return jsonify({"summary": summary})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/summarize_file', methods=['POST'])
def summarize_file_route():
    try:
        file = request.files.get('file')
        if not file:
            return "Error: No file uploaded!", 400
        
        file_ext = os.path.splitext(file.filename)[-1].lower()
        temp_path = f"temp{file_ext}"
        file.save(temp_path)

        if file_ext == ".pdf":
            text = extract_text_from_pdf(temp_path)
        elif file_ext == ".docx":
            text = extract_text_from_docx(temp_path)
        else:
            os.remove(temp_path)
            return "Error: Unsupported file format!", 400

        os.remove(temp_path)  # Clean up temp file
        summary = summarize_text(text)

        # Save Summary to PDF
        summary_pdf_path = "summary.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), summary, fontsize=12)
        doc.save(summary_pdf_path)
        doc.close()

        return send_file(summary_pdf_path, as_attachment=True)
    except Exception as e:
        return f"Error: {str(e)}", 500




# SENTIMENT ANALYSIS TOOL
import torch

# Set model paths
binary_model_path = r"C:\Users\nihar\Downloads\binary_sentiment_model"
emotion_model_path = r"C:\Users\nihar\Downloads\sentiment_model"

# Check for CUDA availability
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(f"Using device: {device}")

# Load binary sentiment model
binary_tokenizer = AutoTokenizer.from_pretrained(binary_model_path)
binary_model = AutoModelForSequenceClassification.from_pretrained(binary_model_path)
binary_model.to(device)

print("Binary sentiment model loaded to", device)

# Load emotion classification model
emotion_tokenizer = AutoTokenizer.from_pretrained(emotion_model_path)
emotion_model = AutoModelForSequenceClassification.from_pretrained(emotion_model_path)
emotion_model.to(device)

print("Emotion classification model loaded to", device)


# Labels
binary_labels = ["Negative", "Positive"]
emotion_labels = ["joy","fear","anger","love","sad","suprise"]

def predict_sentiment(texts, tokenizer, model, labels):
    # device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    
    # Tokenize and move inputs to the correct device
    inputs = tokenizer(texts, padding=True, truncation=True, return_tensors="pt", max_length=512)
    inputs = {k: v.to(device) for k, v in inputs.items()}
    
    model.to(device)
    model.eval()

    with torch.no_grad():
        outputs = model(**inputs)
        preds = torch.argmax(outputs.logits, dim=1).cpu().numpy()  # move to CPU before converting to numpy

    return [labels[p] for p in preds]

from flask import Flask, render_template, request, jsonify, send_file
import pandas as pd
from io import BytesIO

# Assume binary_model, emotion_model, binary_tokenizer, emotion_tokenizer, etc. are already loaded

@app.route('/sentiment', methods=["GET", "POST"])
def sentiment():
    if request.method == "POST":
        csv_file = request.files.get('csv_file')
        text_column = request.form.get('text_column')

        if not csv_file:
            return "No CSV file uploaded", 400
        if not text_column:
            return "Text column not specified", 400

        encodings_to_try = ["utf-8", "ISO-8859-1", "cp1252"]
        df = None

        for enc in encodings_to_try:
            try:
                csv_file.seek(0)
                df = pd.read_csv(csv_file, encoding=enc, nrows=1000)
                break
            except Exception:
                continue

        if df is None:
            return "Unable to read CSV file with common encodings", 400

        if text_column not in df.columns:
            return f"Column '{text_column}' not found in CSV", 400

        print("Available columns:", df.columns.tolist())
        print("Selected column:", text_column)

        texts = df[text_column].astype(str).fillna("").tolist()

        binary_preds = predict_sentiment(texts, binary_tokenizer, binary_model, binary_labels)
        emotion_preds = predict_sentiment(texts, emotion_tokenizer, emotion_model, emotion_labels)

        df['Binary Sentiment'] = binary_preds
        df['Emotion'] = emotion_preds

        pdf_buffer = generate_sentiment_report(df, text_column)

        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name="sentiment_report.pdf",
            mimetype='application/pdf'
        )

    return render_template("sentiment.html")


@app.route("/get_csv_columns", methods=["POST"])
def get_csv_columns():
    try:
        file = request.files.get("csv_file")
        if not file:
            return jsonify({"error": "No file uploaded"}), 400

        encodings_to_try = ["utf-8", "ISO-8859-1", "cp1252"]
        df = None

        for enc in encodings_to_try:
            try:
                file.seek(0)
                df = pd.read_csv(file, encoding=enc)
                break
            except Exception:
                continue

        if df is None:
            return jsonify({"error": "Unable to read CSV with common encodings"}), 400

        return jsonify(df.columns.tolist())
    except Exception as e:
        return jsonify({"error": str(e)}), 500




def generate_sentiment_report(df, text_column):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", "B", 18)
    pdf.set_text_color(0, 0, 128)
    pdf.cell(200, 10, "Sentiment Analysis Report", ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Arial", size=12)
    pdf.ln(8)

    # --- Emotion Distribution (Bar)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Emotion Distribution - Bar Chart", ln=True)
    emotion_counts = df['Emotion'].value_counts()
    fig1, ax1 = plt.subplots()
    sns.barplot(x=emotion_counts.index, y=emotion_counts.values, ax=ax1)
    ax1.set_title("Emotion Distribution")
    plt.xticks(rotation=45)
    img1 = save_plot_to_buffer(fig1)
    pdf.image(img1, x=15, y=None, w=180)
    pdf.ln(10)

    # --- Emotion Distribution (Pie)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Emotion Distribution - Pie Chart", ln=True)
    fig2, ax2 = plt.subplots()
    ax2.pie(emotion_counts.values, labels=emotion_counts.index, autopct='%1.1f%%', startangle=90)
    ax2.set_title("Emotion Proportions")
    img2 = save_plot_to_buffer(fig2)
    pdf.image(img2, x=30, y=None, w=150)
    pdf.ln(10)

    # --- Binary Sentiment (Bar)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Binary Sentiment Distribution - Bar Chart", ln=True)
    binary_counts = df['Binary Sentiment'].value_counts()
    fig3, ax3 = plt.subplots()
    sns.barplot(x=binary_counts.index, y=binary_counts.values, ax=ax3)
    ax3.set_title("Binary Sentiment Distribution")
    img3 = save_plot_to_buffer(fig3)
    pdf.image(img3, x=15, y=None, w=180)
    pdf.ln(10)

    # --- Binary Sentiment (Pie)
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Binary Sentiment - Pie Chart", ln=True)
    fig4, ax4 = plt.subplots()
    ax4.pie(binary_counts.values, labels=binary_counts.index, autopct='%1.1f%%', colors=["#FF4444", "#00C49F"], startangle=90)
    ax4.set_title("Sentiment Proportions")
    img4 = save_plot_to_buffer(fig4)
    pdf.image(img4, x=30, y=None, w=150)
    pdf.ln(10)

    # --- Word Cloud: Positive
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Word Cloud - Positive Sentiment", ln=True)
    positive_text = " ".join(df[df['Binary Sentiment'] == 'Positive'][text_column].astype(str))
    wc_pos = WordCloud(width=600, height=300, background_color="white").generate(positive_text)
    fig5, ax5 = plt.subplots()
    ax5.imshow(wc_pos, interpolation='bilinear')
    ax5.axis('off')
    img5 = save_plot_to_buffer(fig5)
    pdf.image(img5, x=10, y=None, w=180)
    pdf.ln(10)

    # --- Word Cloud: Negative
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Word Cloud - Negative Sentiment", ln=True)
    negative_text = " ".join(df[df['Binary Sentiment'] == 'Negative'][text_column].astype(str))
    wc_neg = WordCloud(width=600, height=300, background_color="white").generate(negative_text)
    fig6, ax6 = plt.subplots()
    ax6.imshow(wc_neg, interpolation='bilinear')
    ax6.axis('off')
    img6 = save_plot_to_buffer(fig6)
    pdf.image(img6, x=10, y=None, w=180)
    pdf.ln(10)

    # --- Stacked Bar: Emotion vs Sentiment
    pdf.set_font("Arial", "B", 14)
    pdf.cell(0, 10, "Emotion Distribution by Sentiment", ln=True)
    emotion_sentiment_counts = df.groupby(['Emotion', 'Binary Sentiment']).size().unstack().fillna(0)
    fig7, ax7 = plt.subplots()
    emotion_sentiment_counts.plot(kind='bar', stacked=True, ax=ax7, colormap="coolwarm")
    ax7.set_title("Emotion vs Sentiment")
    ax7.set_ylabel("Count")
    plt.xticks(rotation=45)
    img7 = save_plot_to_buffer(fig7)
    pdf.image(img7, x=10, y=None, w=180)
    pdf.ln(10)

    # Output to buffer
    pdf_buffer = io.BytesIO()
    pdf_data = pdf.output(dest='S').encode('latin1')
    pdf_buffer.write(pdf_data)
    pdf_buffer.seek(0)
    return pdf_buffer


def save_plot_to_buffer(fig):
    tmpfile = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmpfile.name, bbox_inches='tight')
    plt.close(fig)
    return tmpfile.name








if __name__ == '__main__':
    app.run(debug=True)